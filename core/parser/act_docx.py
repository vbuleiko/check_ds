"""Парсер актов приёмки выполненных работ (docx)."""
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document


MONTHS_RU = {
    'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
    'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
    'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12,
}

MONTH_NAMES_RU = {
    1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель',
    5: 'Май', 6: 'Июнь', 7: 'Июль', 8: 'Август',
    9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь',
}


@dataclass
class ActData:
    filename: str
    contract_full_number: str
    contract_number: str        # "219", "220", "222", "252"
    period_year: int
    period_month: int
    total_km_fact: float        # Фактически км (строка Всего для км)
    total_price: float          # Сумма к оплате (п.4 акта), руб.
    parse_errors: list = field(default_factory=list)

    @property
    def period_name(self) -> str:
        return MONTH_NAMES_RU.get(self.period_month, str(self.period_month))


def _clean_number(text: str) -> Optional[float]:
    """Парсит число вида '1 648 223,75' или '1\xa0648\xa0223,75' → float."""
    text = text.strip().replace('\xa0', '').replace('\u2009', '').replace(' ', '').replace(',', '.')
    try:
        return float(text)
    except ValueError:
        return None


def parse_act_docx(file_path: Path) -> ActData:
    errors: list[str] = []

    doc = Document(str(file_path))

    contract_full_number = ''
    contract_number = ''
    period_year = 0
    period_month = 0
    total_km_fact = 0.0
    total_price = 0.0

    # --- Парсим параграфы ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Полный номер контракта: "№ 01722000025210002220001 от..."
        if not contract_full_number:
            m = re.search(r'№\s+(\d{20,25})', text)
            if m:
                full = m.group(1)
                contract_full_number = full
                # Контракт: символы [-7:-4] = трёхзначный номер (219/220/222/252)
                if len(full) >= 7:
                    contract_number = full[-7:-4]

        # Период: "за период с 01 января 2026 г. по 31 января 2026 г."
        if not period_month:
            m = re.search(r'за период с \d+\s+(\w+)\s+(\d{4})', text, re.IGNORECASE)
            if m:
                month_name = m.group(1).lower()
                period_month = MONTHS_RU.get(month_name, 0)
                period_year = int(m.group(2))

        # Сумма к оплате (п.4): "подлежащая оплате за фактически выполненный объем работ, составляет"
        if not total_price and 'подлежащая оплате за фактически выполненный объем работ' in text:
            # Число перед первой скобкой "("
            m = re.search(r'составляет[\s\xa0]+([\d\s\xa0\u2009]+?)\s*\(', text)
            if m:
                rubles = _clean_number(m.group(1))
                if rubles is not None:
                    # Копейки: "рублей XX копеек"
                    km = re.search(r'\)\s*рубл\w*[\s\xa0]+(\d+)[\s\xa0]*копе', text)
                    kopecks = int(km.group(1)) if km else 0
                    total_price = round(rubles + kopecks / 100, 2)

    # --- Парсим таблицу объёмов (таблица с индексом 1) ---
    for i, table in enumerate(doc.tables):
        if i != 1:
            continue
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 6:
                continue
            # Строка: "Километры пробега транспортных средств" | "км" | "Всего" | ... | факт_значение
            is_km_row = 'Километры пробега транспортных средств' in cells[1]
            is_vsego = cells[3] == 'Всего'
            is_km_unit = cells[2] == 'км'
            if is_km_row and is_km_unit and is_vsego:
                val = _clean_number(cells[5])
                if val is not None:
                    total_km_fact = val
                break

    # --- Проверяем результаты ---
    if not contract_number:
        errors.append('Не удалось определить номер контракта')
    if not period_month:
        errors.append('Не удалось определить период')
    if total_km_fact == 0.0:
        errors.append('Не удалось извлечь объём км (строка "Всего")')
    if total_price == 0.0:
        errors.append('Не удалось извлечь сумму к оплате (п.4 акта)')

    return ActData(
        filename=file_path.name,
        contract_full_number=contract_full_number,
        contract_number=contract_number,
        period_year=period_year,
        period_month=period_month,
        total_km_fact=total_km_fact,
        total_price=total_price,
        parse_errors=errors,
    )
