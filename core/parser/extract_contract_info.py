#!/usr/bin/env python3
"""
Скрипт для извлечения информации из архива с дополнительным соглашением.
Поддерживает ZIP и RAR архивы.
"""

import json
import re
import shutil
import subprocess
import sys
import zipfile
import tempfile
import os
from pathlib import Path
from io import BytesIO

try:
    import rarfile
    RAR_SUPPORT = True
except ImportError:
    RAR_SUPPORT = False

# Путь к 7z
SEVENZIP_PATHS = [
    "C:/Program Files/7-Zip/7z.exe",
    "C:/Program Files (x86)/7-Zip/7z.exe",
    "7z",
]

def find_7z():
    """Находит путь к 7z.exe."""
    for path in SEVENZIP_PATHS:
        if Path(path).exists() or shutil.which(path):
            return path
    return None

def extract_rar_with_7z(archive_path: Path, dest_dir: Path) -> bool:
    """Распаковывает RAR архив с помощью 7z."""
    sevenzip = find_7z()
    if not sevenzip:
        return False

    try:
        result = subprocess.run(
            [sevenzip, "x", "-y", f"-o{dest_dir}", str(archive_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        return result.returncode == 0
    except Exception:
        return False

try:
    from docx import Document
except ImportError:
    print("Требуется установить python-docx: pip install python-docx")
    sys.exit(1)

# Поддержка .doc через win32com (Windows)
try:
    import win32com.client
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

# Импорт парсера для маршрута 315 (XLS формат)
try:
    from core.parser.parse_route_315_xls import parse_route_315_xls
    ROUTE_315_SUPPORT = True
except ImportError:
    ROUTE_315_SUPPORT = False


def get_docx_data(file_bytes: bytes) -> tuple[list[str], list[list[list[str]]], list[tuple]]:
    """
    Извлекает параграфы и таблицы из docx файла.

    Returns:
        (paragraphs, tables, body_sequence)
        - paragraphs: все параграфы документа (включая параграфы внутри таблиц)
        - tables: список таблиц, каждая таблица - список строк, каждая строка - список ячеек
        - body_sequence: элементы тела документа в порядке их следования:
            ('para', text) для параграфов верхнего уровня,
            ('table', index) для таблиц (index — позиция в списке tables)
    """
    from docx.text.paragraph import Paragraph as DocxParagraph

    doc = Document(BytesIO(file_bytes))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables.append(table_data)

    # Строим упорядоченную последовательность тела документа (только верхний уровень)
    body_sequence = []
    table_counter = 0
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = DocxParagraph(child, doc).text.strip()
            if text:
                body_sequence.append(('para', text))
        elif tag == 'tbl':
            body_sequence.append(('table', table_counter))
            table_counter += 1

    return paragraphs, tables, body_sequence


def get_doc_data(file_bytes: bytes) -> tuple[list[str], list[list[list[str]]]]:
    """
    Извлекает параграфы и таблицы из doc файла (старый формат Word).
    """
    if not DOC_SUPPORT:
        raise ImportError("Для работы с .doc файлами установите: pip install pywin32")

    # COM требует STA-инициализацию в каждом потоке (нужно для run_in_executor)
    try:
        import pythoncom
        pythoncom.CoInitialize()
        _com_initialized = True
    except Exception:
        _com_initialized = False

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(tmp_path)

        paragraphs = []
        for para in doc.Paragraphs:
            text = para.Range.Text.strip()
            if text:
                paragraphs.append(text)

        tables = []
        for table in doc.Tables:
            table_data = []
            for row_idx in range(1, table.Rows.Count + 1):
                row_data = []
                for col_idx in range(1, table.Columns.Count + 1):
                    try:
                        cell_text = table.Cell(row_idx, col_idx).Range.Text.strip()
                        # Убираем служебные символы Word
                        cell_text = cell_text.replace("\r", "").replace("\x07", "")
                        row_data.append(cell_text)
                    except Exception:
                        row_data.append("")
                table_data.append(row_data)
            tables.append(table_data)

        doc.Close(False)
        word.Quit()

        return paragraphs, tables
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        if _com_initialized:
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except Exception:
                pass


def get_word_data(file_bytes: bytes, filename: str) -> tuple[list[str], list[list[list[str]]], list[tuple]]:
    """Извлекает данные из Word файла (.doc или .docx).

    Returns:
        (paragraphs, tables, body_sequence)
        body_sequence пуст для .doc файлов (нет поддержки порядка элементов через COM).
    """
    if filename.lower().endswith(".docx"):
        return get_docx_data(file_bytes)
    elif filename.lower().endswith(".doc"):
        paragraphs, tables = get_doc_data(file_bytes)
        return paragraphs, tables, []
    else:
        raise ValueError(f"Неподдерживаемый формат: {filename}")


def decode_filename(name: str) -> str:
    """Пытается декодировать имя файла из разных кодировок."""
    # Пробуем перекодировать из cp437 в cp866 (частая проблема с русскими именами в zip)
    try:
        decoded = name.encode('cp437').decode('cp866')
        return decoded
    except (UnicodeDecodeError, UnicodeEncodeError):
        pass
    return name


def find_main_document(archive_files: list[str]) -> str | None:
    """Находит основной документ ДС в списке файлов архива."""
    for file_path in archive_files:
        # Пробуем декодировать имя файла
        decoded_path = decode_filename(file_path)
        filename = Path(decoded_path).name.lower()
        if filename.startswith("дс") and filename.endswith(".docx"):
            return file_path  # Возвращаем оригинальный путь для чтения
    return None


def find_appendix_folders(archive_files: list[str]) -> dict[str, list[str]]:
    """Находит папки приложений и файлы в них."""
    appendix_folders = {}

    for file_path in archive_files:
        decoded_path = decode_filename(file_path)
        match = re.search(r"[Пп]рил[а-я]*\s*№?\s*(\d+)", decoded_path)
        if match:
            appendix_num = match.group(1)
            if appendix_num not in appendix_folders:
                appendix_folders[appendix_num] = []
            if not file_path.endswith("/"):
                appendix_folders[appendix_num].append(file_path)

    return appendix_folders


def find_contract_doc_in_folder(files: list[str]) -> str | None:
    """Находит файл документа приложения в списке файлов папки."""
    for file_path in files:
        decoded_path = decode_filename(file_path)
        filename = Path(decoded_path).name
        # Пропускаем временные файлы Word (~$...)
        if filename.startswith("~$"):
            continue
        filename_lower = filename.lower()
        if filename_lower.endswith(".doc") or filename_lower.endswith(".docx"):
            return file_path  # Возвращаем оригинальный путь
    return None


def find_xls_file_in_folder(files: list[str]) -> str | None:
    """Находит XLS/XLSX файл в папке (без 'к параметрам' в названии)."""
    for file_path in files:
        decoded_path = decode_filename(file_path)
        filename = Path(decoded_path).name.lower()
        if (filename.endswith(".xls") or filename.endswith(".xlsx")):
            if "к параметрам" not in filename:
                return file_path
    return None


def get_route_for_appendix(appendix_num: str, changes: list[dict]) -> str | None:
    """Определяет маршрут для приложения по списку изменений."""
    for change in changes:
        if change.get("appendix") == appendix_num:
            return change.get("route")
    return None


def convert_date_to_iso(date_str: str) -> str:
    """Преобразует дату из DD.MM.YYYY в YYYY-MM-DD."""
    match = re.match(r"(\d{2})\.(\d{2})\.(\d{4})", date_str)
    if match:
        day, month, year = match.groups()
        return f"{year}-{month}-{day}"
    return date_str


def parse_dates(text: str) -> dict:
    """Извлекает даты из текста."""
    result = {"date_from": None, "date_to": None, "date_on": None}

    # Диапазон: "с DD.MM.YYYY года по DD.MM.YYYY года"
    range_match = re.search(
        r"с\s+(\d{2}\.\d{2}\.\d{4})\s+года\s+по\s+(\d{2}\.\d{2}\.\d{4})\s+года",
        text
    )
    if range_match:
        result["date_from"] = convert_date_to_iso(range_match.group(1))
        result["date_to"] = convert_date_to_iso(range_match.group(2))
        return result

    # Начало: "с DD.MM.YYYY года"
    from_match = re.search(r"с\s+(\d{2}\.\d{2}\.\d{4})\s+года", text)
    if from_match:
        result["date_from"] = convert_date_to_iso(from_match.group(1))
        return result

    # Конкретная дата: "на DD.MM.YYYY года"
    on_match = re.search(r"на\s+(\d{2}\.\d{2}\.\d{4})\s+года", text)
    if on_match:
        result["date_on"] = convert_date_to_iso(on_match.group(1))
        return result

    # Конкретная дата в начале строки: "DD.MM.YYYY года"
    on_match2 = re.match(r"^(\d{2}\.\d{2}\.\d{4})\s+года", text)
    if on_match2:
        result["date_on"] = convert_date_to_iso(on_match2.group(1))
        return result

    return result


def extract_route_120_subitem(text: str) -> list[dict]:
    """
    Извлекает информацию из подпункта для маршрута 120 (контракт 2520001).

    Формат: "1.1. 30.12.2025 года - по графику движения..."
    Или: "с 31.12.2025 года по 11.01.2026 года (включительно) - по графику движения..."
    Может содержать несколько дат через ";"

    Возвращает список записей.
    """
    results = []

    # Убираем номер подпункта в начале (1.1. или 1.2. и т.д.)
    # Используем негативный просмотр вперёд, чтобы не ловить даты (DD.MM.YYYY)
    text = re.sub(r"^\d{1,2}\.\d{1,2}\.(?!\d{4})\s*", "", text)

    # Разбиваем по ";" для обработки нескольких дат в одной строке
    parts = [p.strip() for p in text.split(";") if p.strip()]

    for part in parts:
        result = {
            "appendix": None,
            "route": "120",
            "date_from": None,
            "date_to": None,
            "date_on": None,
        }

        # Ищем point - текст после " - " (график движения)
        point_match = re.search(r"\s+-\s+(.+?)\.?\s*$", part)
        if point_match:
            result["point"] = point_match.group(1).strip().rstrip('.')

        # Паттерн: "с DD.MM.YYYY года по DD.MM.YYYY года"
        range_match = re.search(
            r"с\s+(\d{2}\.\d{2}\.\d{4})\s+года\s+по\s+(\d{2}\.\d{2}\.\d{4})\s+года",
            part
        )
        if range_match:
            result["date_from"] = convert_date_to_iso(range_match.group(1))
            result["date_to"] = convert_date_to_iso(range_match.group(2))
            results.append(result)
            continue

        # Паттерн: "DD.MM.YYYY года - " (конкретная дата в начале)
        on_match = re.match(r"^(\d{2}\.\d{2}\.\d{4})\s+года\s+-", part)
        if on_match:
            result["date_on"] = convert_date_to_iso(on_match.group(1))
            results.append(result)
            continue

        # Если нашли point, но не нашли дату - всё равно добавим
        if result.get("point"):
            results.append(result)

    return results


def extract_route_info(text: str) -> list[dict]:
    """Извлекает информацию о маршрутах из текста подпункта."""
    results = []
    routes = []

    # Быстрая проверка - если нет "маршрут", пропускаем
    if "маршрут" not in text.lower():
        return results

    # Ищем "по маршруту №" или "маршруту №" (без "по") - оба могут содержать список маршрутов
    # Список маршрутов может заканчиваться на " - по ", " по ", " согласно" или конец строки
    route_match = re.search(r"(?:по\s+)?маршрут\w*\s*№\s*([\d\s,А-ЯЁа-яёA-Za-z]{1,100}?)(?:\s*-\s+по\s|\s+по\s|\s+согласно|\s*$)", text, re.IGNORECASE)
    if route_match:
        routes_str = route_match.group(1)
        route_parts = re.findall(r"(\d+[А-ЯЁа-яёA-Za-z]*)", routes_str)
        routes.extend(route_parts)

    appendix_match = re.search(r"Приложению\s*№\s*(\d+)", text)
    appendix_num = appendix_match.group(1) if appendix_match else None

    # Извлекаем point - текст после маршрутов (до "согласно Приложению" или конца)
    point = None
    if not appendix_num and routes:
        # Ищем текст после последнего номера маршрута
        # Находим позицию после списка маршрутов
        last_route = routes[-1]
        last_route_pos = text.rfind(last_route)
        if last_route_pos != -1:
            after_routes = text[last_route_pos + len(last_route):]
            # Ищем "по графику..." или подобное
            point_match = re.search(r"[,\s]*(по\s+.+?)\.?\s*$", after_routes, re.IGNORECASE)
            if point_match:
                point = point_match.group(1).strip().rstrip('.')

    dates = parse_dates(text)

    for route in routes:
        result = {
            "appendix": appendix_num,
            "route": route,
            **dates
        }
        if point is not None:
            result["point"] = point
        results.append(result)

    return results


def parse_money_amount(text: str) -> float | None:
    """Преобразует сумму в число."""
    match = re.search(r"([\d\s]+)\s*рубл\w*\s+(\d+)\s*копе", text, re.IGNORECASE)
    if match:
        rubles = match.group(1).replace(" ", "").replace("\u00a0", "")
        kopeks = match.group(2)
        return float(f"{rubles}.{kopeks}")

    match = re.search(r"([\d\s]+)\s*рубл\w*", text, re.IGNORECASE)
    if match:
        rubles = match.group(1).replace(" ", "").replace("\u00a0", "")
        return float(rubles)

    return None


def extract_values_from_tables(tables: list[list[list[str]]]) -> dict:
    """
    Извлекает значения probeg и sum из разных таблиц документа.

    Возвращает словарь с ключами:
    - probeg_sravnenie: из таблицы "Расчет изменения объема", столбец "Предлагаемые изменения (в км)"
    - probeg_etapy: из таблицы "Этапы исполнения Контракта (в части сроков выполнения работ)",
                    столбец "Максимальная транспортная работа", строка Итого
    - probeg_etapy_avans: из таблицы "Этапы исполнения Контракта (с учетом порядка погашения...)",
                          столбец "Максимальная транспортная работа", строка Итого
    - sum_etapy: из таблицы "Этапы исполнения Контракта (в части сроков выполнения работ)",
                 столбец "Стоимость транспортной работы, руб.", строка Итого
    - sum_etapy_avans: из таблицы "Этапы исполнения Контракта (с учетом порядка погашения...)",
                       столбец "Стоимость транспортной работы по этапу, руб. (с учетом выплаченных авансов)", строка Итого
    - sum_finansirovanie_table: из таблицы со столбцом "Финансирование по годам, руб.", строка Итого
    """
    result = {
        "probeg_sravnenie": None,
        "probeg_etapy": None,
        "probeg_etapy_avans": None,
        "sum_etapy": None,
        "sum_etapy_avans": None,
        "sum_finansirovanie_table": None,
    }

    # Счётчик таблиц с "Максимальная транспортная работа" и строкой ИТОГО
    etapy_table_count = 0

    for table in tables:
        if len(table) < 2:
            continue

        header_row = table[0]

        # 1. Таблица "Расчет изменения объема" -> probeg_sravnenie
        # Определяем по наличию столбца "Предлагаемые изменения (в км)"
        if result["probeg_sravnenie"] is None:
            for col_idx, cell in enumerate(header_row):
                cell_lower = cell.lower()
                if "предлагаемые изменения" in cell_lower and "км" in cell_lower:
                    # Значение в строке 1 (первая строка данных после заголовка)
                    if len(table) > 1 and col_idx < len(table[1]):
                        result["probeg_sravnenie"] = parse_number(table[1][col_idx])
                    break

        # 2. Таблица "Финансирование по годам" -> sum_finansirovanie_table
        if result["sum_finansirovanie_table"] is None:
            for col_idx, cell in enumerate(header_row):
                cell_lower = cell.lower()
                if "финансирование по годам" in cell_lower:
                    # Ищем строку "Итого"
                    for row in table:
                        if row and "итого" in row[0].lower():
                            if col_idx < len(row):
                                result["sum_finansirovanie_table"] = parse_number(row[col_idx])
                            break
                    break

        # 3-4. Таблицы "Этапы исполнения Контракта" -> probeg_etapy/sum_etapy и probeg_etapy_avans/sum_etapy_avans
        # Определяем по наличию столбца "Максимальная транспортная работа" и строки ИТОГО
        # Пропускаем таблицу "Расчет изменения объема" (она имеет столбец "Предлагаемые изменения")
        is_sravnenie_table = any("предлагаемые изменения" in cell.lower() for cell in header_row)
        if is_sravnenie_table:
            continue

        probeg_col_idx = None
        sum_col_idx = None
        sum_avans_col_idx = None

        for col_idx, cell in enumerate(header_row):
            cell_lower = cell.lower()
            if "максимальная транспортная работа" in cell_lower:
                probeg_col_idx = col_idx
            if "стоимость транспортной работы" in cell_lower and "руб" in cell_lower:
                if "с учетом выплаченных авансов" in cell_lower or "с учётом выплаченных авансов" in cell_lower:
                    sum_avans_col_idx = col_idx
                elif sum_col_idx is None:
                    sum_col_idx = col_idx

        # Таблица этапов должна содержать И пробег И стоимость
        if probeg_col_idx is not None and sum_col_idx is not None:
            # Ищем строку "Итого"
            for row in table:
                if row and "итого" in row[0].lower():
                    etapy_table_count += 1
                    if etapy_table_count == 1:
                        # Первая таблица -> probeg_etapy, sum_etapy
                        if probeg_col_idx < len(row):
                            result["probeg_etapy"] = parse_number(row[probeg_col_idx])
                        if sum_col_idx < len(row):
                            result["sum_etapy"] = parse_number(row[sum_col_idx])
                    elif etapy_table_count == 2:
                        # Вторая таблица -> probeg_etapy_avans, sum_etapy_avans
                        if probeg_col_idx < len(row):
                            result["probeg_etapy_avans"] = parse_number(row[probeg_col_idx])
                        if sum_avans_col_idx is not None and sum_avans_col_idx < len(row):
                            result["sum_etapy_avans"] = parse_number(row[sum_avans_col_idx])
                    break

    return result


def extract_stages_km_252(tables: list[list[list[str]]]) -> dict | None:
    """
    Извлекает данные км по периодам из таблицы ГК252.

    Находит таблицу со столбцом "Расчетный период" и "Максимальная транспортная работа"
    но без "Стоимость транспортной работы". Это таблица этапов только с пробегом.

    Returns:
        dict совместимый с KmData.to_dict() или None если таблица не найдена
    """
    MONTH_NAMES_252 = {
        "январь": 1, "февраль": 2, "март": 3, "апрель": 4,
        "май": 5, "июнь": 6, "июль": 7, "август": 8,
        "сентябрь": 9, "октябрь": 10, "ноябрь": 11, "декабрь": 12,
    }

    for table in tables:
        if len(table) < 5:
            continue

        header = table[0]
        has_km = any("максимальная транспортная работа" in str(c).lower() for c in header)
        has_raschetny = any("расчетный период" in str(c).lower() for c in header)
        has_stoimost = any("стоимость транспортной работы" in str(c).lower() for c in header)

        if not (has_km and has_raschetny and not has_stoimost):
            continue

        # Находим индекс столбца "Максимальная транспортная работа"
        km_col = None
        for i, c in enumerate(header):
            if "максимальная транспортная работа" in str(c).lower():
                km_col = i
                break
        if km_col is None:
            continue

        total_col = km_col + 1  # столбец "Итого"

        monthly = []
        grand_total = None

        # Данные начинаются с row 3 (skip header rows 0, 1, 2)
        for row in table[3:]:
            if not row or len(row) < 4:
                continue

            row0 = str(row[0]).strip().lower() if row[0] else ""

            # Строка ИТОГО
            if "итого" in row0:
                val_str = str(row[total_col]).strip() if total_col < len(row) else ""
                if not val_str or val_str.lower() == "итого":
                    val_str = str(row[km_col]).strip() if km_col < len(row) else ""
                if val_str and val_str.lower() != "итого":
                    grand_total = parse_number(val_str)
                break

            # Год (col 1)
            year_str = str(row[1]).strip() if len(row) > 1 else ""
            year_num = parse_number(year_str)
            if not year_num or not (2020 <= int(year_num) <= 2035):
                continue
            year = int(year_num)

            # Квартал/месяц (col 2)
            period_label = str(row[2]).strip() if len(row) > 2 else ""

            # Расчетный период (col 3): "15.07-30.09" или "01.01-31.01"
            calc_period = str(row[3]).strip() if len(row) > 3 else ""
            pm = re.match(r'^(\d+)\.(\d+)-(\d+)\.(\d+)$', calc_period)
            if not pm:
                continue

            start_day = int(pm.group(1))
            start_month = int(pm.group(2))
            end_day = int(pm.group(3))
            end_month = int(pm.group(4))

            # Км (col km_col)
            km_str = str(row[km_col]).strip() if km_col < len(row) else ""
            km_val = parse_number(km_str)
            if km_val is None or float(km_val) <= 0:
                continue
            km_val = round(float(km_val), 2)

            # Месяц или квартал?
            month_num = MONTH_NAMES_252.get(period_label.lower().strip())

            if month_num is not None:
                # Месячный период
                monthly.append({
                    "year": year,
                    "month": month_num,
                    "period_start": start_day,
                    "period_end": end_day,
                    "routes": {"120": km_val},
                    "total": km_val,
                    "month_end": None,
                })
            else:
                # Квартальный период (label = "1", "2", "3", "4")
                monthly.append({
                    "year": year,
                    "month": start_month,
                    "period_start": start_day,
                    "period_end": end_day,
                    "routes": {"120": km_val},
                    "total": km_val,
                    "month_end": end_month,
                })

        if monthly:
            return {
                "contract": "252",
                "ds_number": None,
                "grand_total": grand_total,
                "monthly": monthly,
            }

    return None


def detect_vysvobozhdenie(paragraphs: list[str], tables: list[list[list[str]]]) -> dict | None:
    """
    Проверяет, содержит ли документ ДС информацию о высвобождении.
    
    Признаки высвобождения:
    - "стоимость фактически выполненных работ за N этап составила"
    - "цена Контракта составит" (новая цена после высвобождения)
    
    Returns:
        dict с данными высвобождения или None если не найдено
    """
    text = " ".join(paragraphs)
    
    # Проверяем наличие ключевых фраз
    # Используем .{0,100} чтобы пропустить вставные слова вида
    # "Подрядчиком и принятых Заказчиком по Контракту" между "выполненных" и "работ"
    has_fact_work = re.search(
        r"стоимость\s+фактически\s+выполненных.{0,100}работ\s+за\s+(\d+)\s+этап",
        text,
        re.IGNORECASE | re.DOTALL
    )
    
    if not has_fact_work:
        return None
    
    # Номер закрываемого этапа
    closed_stage = int(has_fact_work.group(1)) if has_fact_work else None
    
    # Сумма фактически выполненных работ
    fact_amount_match = re.search(
        r"составила\s+([\d\s\u00a0]+\s*руб\w*(?:\s+\d+\s*коп\w*)?)",
        text,
        re.IGNORECASE
    )
    closed_amount = parse_money_amount(fact_amount_match.group(1)) if fact_amount_match else None
    
    # Новая цена контракта
    new_price_match = re.search(
        r"цена\s+[Кк]онтракта\s+составит\s+([\d\s\u00a0]+\s*руб\w*(?:\s+\d+\s*коп\w*)?)",
        text,
        re.IGNORECASE
    )
    new_contract_price = parse_money_amount(new_price_match.group(1)) if new_price_match else None

    # Парсим таблицы этапов из высвобождения (аналогично vysvobozhdenie_parser)
    stages_data = _extract_vysvobozhdenie_tables(tables)

    # Если цена не найдена в тексте — используем ИТОГО из таблицы этапов
    if new_contract_price is None:
        new_contract_price = stages_data.get("itogo_price")

    return {
        "closed_stage": closed_stage,
        "closed_amount": closed_amount,
        "new_contract_price": new_contract_price,
        **stages_data
    }


def _extract_vysvobozhdenie_tables(tables: list[list[list[str]]]) -> dict:
    """
    Извлекает таблицы этапов из раздела о высвобождении.
    
    Возвращает данные таблиц для проверки.
    """
    result = {
        "stages_table1_raw": None,
        "stages_table2_raw": None,
        "stages_finansirovanie_raw": None,
        "itogo_km": None,
        "itogo_price": None,
    }
    
    for table in tables:
        if not table or len(table) < 2:
            continue
            
        header = table[0]
        header_text = " ".join(header).lower()
        
        # Финансирование по годам
        if result["stages_finansirovanie_raw"] is None and "финансирование" in header_text and "транспортная" not in header_text:
            result["stages_finansirovanie_raw"] = table
            continue
        
        # Таблицы этапов
        if "максимальная транспортная работа" in header_text:
            has_price = "стоимость" in header_text
            has_avans = "аванс" in header_text
            
            if result["stages_table2_raw"] is None and has_avans:
                result["stages_table2_raw"] = table
                continue
            
            if result["stages_table1_raw"] is None and has_price and not has_avans:
                result["stages_table1_raw"] = table
                continue
    
    # Извлекаем итоги из таблицы 1
    if result["stages_table1_raw"]:
        for row in result["stages_table1_raw"]:
            if row and "итого" in str(row[0]).lower():
                # Находим колонки км и цены
                km_col = None
                price_col = None
                for idx, cell in enumerate(result["stages_table1_raw"][0]):
                    cell_lower = cell.lower()
                    if "максимальная транспортная работа" in cell_lower:
                        km_col = idx
                    elif "стоимость транспортной работы" in cell_lower and "аванс" not in cell_lower:
                        price_col = idx
                
                if km_col is not None and len(row) > km_col:
                    result["itogo_km"] = parse_number(row[km_col])
                if price_col is not None and len(row) > price_col:
                    result["itogo_price"] = parse_number(row[price_col])
                break
    
    return result


def extract_general_info(paragraphs: list[str]) -> dict:
    """Извлекает общую информацию: номер ДС, контракта и сумму."""
    text = "\n".join(paragraphs)
    result = {
        "ds_number": None,
        "contract_number": None,
        "sum_text": None,
        "sum_finansirovanie_text": None,
        "sum_etapy": None,
        "sum_etapy_avans": None,
        "sum_finansirovanie_table": None,
        "probeg_sravnenie": None,
        "probeg_etapy": None,
        "probeg_etapy_avans": None,
        "price_change_direction": None,
        "price_change_amount": None,
    }

    ds_patterns = [
        r"[Дд]ополнительное\s+соглашение\s*[№#N]\s*(\d+)",
        r"ДС\s*[№#N]\s*(\d+)",
    ]
    for pattern in ds_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["ds_number"] = match.group(1)
            break

    contract_patterns = [
        r"[Гг]осударственн\w*\s+контракт\w*\s*[№#N]\s*(\d{10,})",
        r"контракт\w*\s*[№#N]\s*(\d{10,})",
    ]
    for pattern in contract_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["contract_number"] = match.group(1)
            break

    # sum_text - из текста "изменить цену Контракта...составляет: X рублей Y копеек"
    sum_match = re.search(
        r"изменить\s+цену\s+[Кк]онтракта.*?составляет[:\s]*([\d\s]+\s*рубл\w*\s+\d+\s*копе\w*)",
        text,
        re.IGNORECASE | re.DOTALL
    )
    if sum_match:
        result["sum_text"] = parse_money_amount(sum_match.group(1))

    # price_change_direction и price_change_amount — из текста
    # "Стороны пришли к соглашению увеличить/уменьшить цену Контракта на X рублей Y копеек"
    change_match = re.search(
        r"(увеличить|уменьшить)\s+цену\s+[Кк]онтракта\s+на\s+"
        r"([\d\s\u00a0]+\s*рубл\w*(?:\s+\d+\s*копе\w*)?)",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if change_match:
        result["price_change_direction"] = change_match.group(1).lower()
        result["price_change_amount"] = parse_money_amount(change_match.group(2))

    # sum_finansirovanie_text - из текста "ИТОГО: X рублей Y копеек"
    itogo_match = re.search(
        r"ИТОГО:\s*([\d\s]+\s*рубл\w*\s+\d+\s*копе\w*)",
        text,
        re.IGNORECASE
    )
    if itogo_match:
        result["sum_finansirovanie_text"] = parse_money_amount(itogo_match.group(1))

    return result


def extract_changes(paragraphs: list[str], contract_number: str = None) -> tuple[list, list, list, list]:
    """Извлекает данные об изменениях из пунктов документа."""
    change_with_money = []
    change_with_money_no_appendix = []
    change_without_money = []
    change_without_money_no_appendix = []

    section1_start = None
    section1_end = None
    section2_start = None
    section2_end = None
    route_120_section_start = None
    route_120_section_end = None

    # Специальная обработка для контракта 2520001 (маршрут 120)
    is_route_120_contract = contract_number and contract_number.endswith("2520001")

    for i, text in enumerate(paragraphs):
        if "Стороны пришли к соглашению выполнять работы:" in text and section1_start is None:
            section1_start = i
        elif "Стороны пришли к соглашению выполнять работы без изменения" in text and section2_start is None:
            section2_start = i
        # Специальный случай для маршрута 120
        elif is_route_120_contract and "Стороны пришли к соглашению выполнять работы по маршруту № 120" in text:
            route_120_section_start = i

    # Определяем конец секции 1
    if section1_start is not None:
        if section2_start is not None:
            # Есть секция 2 - она начинается после секции 1
            section1_end = section2_start
        else:
            # Нет секции 2 - ищем конец по следующему "Стороны пришли к соглашению"
            for i in range(section1_start + 1, len(paragraphs)):
                if "Стороны пришли к соглашению" in paragraphs[i]:
                    section1_end = i
                    break
            if section1_end is None:
                section1_end = len(paragraphs)

    # Определяем конец секции 2
    if section2_start is not None:
        for i in range(section2_start + 1, len(paragraphs)):
            if paragraphs[i].startswith("Стороны"):
                section2_end = i
                break
        if section2_end is None:
            section2_end = len(paragraphs)

    # Определяем конец секции маршрута 120
    if route_120_section_start is not None:
        for i in range(route_120_section_start + 1, len(paragraphs)):
            # Секция заканчивается на следующем пункте верхнего уровня (не подпункт)
            if re.match(r"^\d+\.\s+[А-ЯЁA-Z]", paragraphs[i]):
                route_120_section_end = i
                break
        if route_120_section_end is None:
            route_120_section_end = len(paragraphs)

    # Извлекаем данные из секции 1 (с изменением цены)
    if section1_start is not None and section1_end is not None:
        for i in range(section1_start + 1, section1_end):
            route_infos = extract_route_info(paragraphs[i])
            for info in route_infos:
                if info.get("appendix"):
                    change_with_money.append(info)
                else:
                    change_with_money_no_appendix.append(info)

    # Извлекаем данные из секции 2 (без изменения цены)
    if section2_start is not None:
        end = section2_end if section2_end else len(paragraphs)
        for i in range(section2_start + 1, end):
            route_infos = extract_route_info(paragraphs[i])
            for info in route_infos:
                if info.get("appendix"):
                    change_without_money.append(info)
                else:
                    change_without_money_no_appendix.append(info)

    # Извлекаем данные из секции маршрута 120 (специальный случай)
    if route_120_section_start is not None and route_120_section_end is not None:
        for i in range(route_120_section_start + 1, route_120_section_end):
            text = paragraphs[i]
            # Пропускаем пустые строки
            if not text.strip():
                continue
            # Обрабатываем строки с датами и "по графику движения"
            # Формат: "1.1. DD.MM.YYYY года - по графику..." или "DD.MM.YYYY года - по графику..."
            if "по графику движения" in text.lower() and re.search(r"\d{2}\.\d{2}\.\d{4}", text):
                infos = extract_route_120_subitem(text)
                change_with_money_no_appendix.extend(infos)

    return change_with_money, change_with_money_no_appendix, change_without_money, change_without_money_no_appendix


def parse_number(val_str: str) -> float | int | None:
    """Парсит число из строки."""
    val_str = val_str.replace(" ", "").replace("\u00a0", "").replace(",", ".")
    try:
        num = float(val_str)
        # Если целое число - возвращаем int
        if num == int(num):
            return int(num)
        return num
    except ValueError:
        return None


def extract_single_trips_table(table: list[list[str]]) -> dict | None:
    """
    Извлекает данные о рейсах из одной таблицы.
    Возвращает dict с num_of_types, type_X_name, type_X_np и т.д. или None.
    """
    if len(table) < 3:
        return None

    # Ищем строку с "Направление" и следующую с "Количество"
    header_row_idx = None
    for idx, row in enumerate(table):
        first_cell = row[0].lower() if row else ""
        if "направлен" in first_cell:
            if idx + 1 < len(table):
                next_row_text = " ".join(table[idx + 1]).lower()
                if "количество" in next_row_text:
                    header_row_idx = idx
                    break

    if header_row_idx is None:
        return None

    # Заголовки типов дней
    header_row = table[header_row_idx]
    day_types = []

    for cell in header_row[1:]:
        cell_clean = cell.strip()
        if cell_clean and "направлен" not in cell_clean.lower():
            if cell_clean not in day_types:
                day_types.append(cell_clean)

    if not day_types:
        return None

    result = {"num_of_types": len(day_types)}

    for i, day_type in enumerate(day_types, 1):
        result[f"type_{i}_name"] = day_type

    # Ищем строки "Прямое", "Обратное" и "ИТОГО"
    np_row = None
    op_row = None
    sum_row = None

    for row in table[header_row_idx + 2:]:
        first_cell = row[0].lower() if row else ""
        if "прямое" in first_cell or "прямо" in first_cell:
            np_row = row
        elif "обратное" in first_cell or "обратн" in first_cell:
            op_row = row
        elif "итого" in first_cell:
            sum_row = row

    if not np_row and not op_row:
        return None

    def extract_row_values(row, num_types):
        trips = []
        probeg = []
        for i in range(num_types):
            trips_col = 1 + i * 2
            probeg_col = 2 + i * 2

            if trips_col < len(row):
                trips.append(parse_number(row[trips_col]))
            else:
                trips.append(None)

            if probeg_col < len(row):
                probeg.append(parse_number(row[probeg_col]))
            else:
                probeg.append(None)

        return trips, probeg

    if np_row:
        np_trips, np_probeg = extract_row_values(np_row, len(day_types))
        for i, val in enumerate(np_trips, 1):
            result[f"type_{i}_forward_number"] = val
        for i, val in enumerate(np_probeg, 1):
            result[f"type_{i}_forward_probeg"] = val

    if op_row:
        op_trips, op_probeg = extract_row_values(op_row, len(day_types))
        for i, val in enumerate(op_trips, 1):
            result[f"type_{i}_reverse_number"] = val
        for i, val in enumerate(op_probeg, 1):
            result[f"type_{i}_reverse_probeg"] = val

    if sum_row:
        sum_trips, sum_probeg = extract_row_values(sum_row, len(day_types))
        for i, val in enumerate(sum_trips, 1):
            result[f"type_{i}_sum_number"] = val
        for i, val in enumerate(sum_probeg, 1):
            result[f"type_{i}_sum_probeg"] = val

    return result


def extract_period_dates(text: str) -> dict:
    """
    Извлекает даты периода из строки вида:
    "с 16.11 по 14.04" -> date_from="0000-11-16", date_to="0000-04-14"
    "с 16.11.2025 по 14.04.2026" -> date_from="2025-11-16", date_to="2026-04-14"
    """
    result = {"date_from": None, "date_to": None, "date_on": None}

    # Паттерн с годом: DD.MM.YYYY
    match_with_year = re.search(
        r"с\s+(\d{1,2})\.(\d{1,2})\.(\d{4})\s+по\s+(\d{1,2})\.(\d{1,2})\.(\d{4})",
        text
    )
    if match_with_year:
        d1, m1, y1, d2, m2, y2 = match_with_year.groups()
        result["date_from"] = f"{y1}-{m1.zfill(2)}-{d1.zfill(2)}"
        result["date_to"] = f"{y2}-{m2.zfill(2)}-{d2.zfill(2)}"
        return result

    # Паттерн без года: DD.MM
    match_no_year = re.search(
        r"с\s+(\d{1,2})\.(\d{1,2})\s+по\s+(\d{1,2})\.(\d{1,2})",
        text
    )
    if match_no_year:
        d1, m1, d2, m2 = match_no_year.groups()
        result["date_from"] = f"0000-{m1.zfill(2)}-{d1.zfill(2)}"
        result["date_to"] = f"0000-{m2.zfill(2)}-{d2.zfill(2)}"
        return result

    return result


def extract_trips_from_tables(tables: list[list[list[str]]], paragraphs: list[str] = None) -> dict:
    """
    Извлекает информацию о рейсах и пробеге из таблиц.

    Если в документе есть таблицы для разных сезонов (зима/лето), возвращает:
    {
        "period_winter": { "date_from": ..., "date_to": ..., "num_of_types": ..., ... },
        "period_summer": { "date_from": ..., "date_to": ..., "num_of_types": ..., ... }
    }

    Иначе возвращает данные напрямую (без обёртки period_*).
    """
    result = {"num_of_types": 0}

    # Проверяем, есть ли в параграфах указания на сезонность
    has_winter = False
    has_summer = False
    winter_dates = {"date_from": None, "date_to": None, "date_on": None}
    summer_dates = {"date_from": None, "date_to": None, "date_on": None}

    if paragraphs:
        for para in paragraphs:
            para_lower = para.lower()
            # Ищем строку с "количество рейсов" и "зима"
            if "количество" in para_lower and "рейсов" in para_lower and "зима" in para_lower:
                has_winter = True
                winter_dates = extract_period_dates(para)
            # Ищем строку с "количество рейсов" и "лето"
            if "количество" in para_lower and "рейсов" in para_lower and "лето" in para_lower:
                has_summer = True
                summer_dates = extract_period_dates(para)

    # Если есть оба сезона - ищем две таблицы
    if has_winter and has_summer:
        winter_data = None
        summer_data = None

        # Собираем все таблицы с данными о рейсах
        trips_tables = []
        for table in tables:
            table_data = extract_single_trips_table(table)
            if table_data and table_data.get("num_of_types", 0) > 0:
                trips_tables.append(table_data)

        # Если нашли ровно 2 таблицы - первая зима, вторая лето
        if len(trips_tables) >= 2:
            winter_data = {**winter_dates, **trips_tables[0]}
            summer_data = {**summer_dates, **trips_tables[1]}
        elif len(trips_tables) == 1:
            # Если только одна таблица, но есть указания на сезоны - кладём в зиму
            winter_data = {**winter_dates, **trips_tables[0]}

        result = {}
        if winter_data:
            result["period_winter"] = winter_data
        if summer_data:
            result["period_summer"] = summer_data

        return result

    # Иначе - старая логика: берём первую подходящую таблицу
    for table in tables:
        table_data = extract_single_trips_table(table)
        if table_data and table_data.get("num_of_types", 0) > 0:
            return table_data

    return result


def extract_appendix_info(paragraphs: list[str], tables: list[list[list[str]]]) -> dict:
    """Извлекает информацию из документа приложения."""
    result = {
        "appendix_num": None,
        "ds_num": None,
        "route": None,
        "date_from": None,
        "date_to": None,
        "date_on": None,
        "length_forward": None,
        "length_reverse": None,
        "length_sum": None
    }

    # Ищем номер приложения и номер ДС: "Приложение №5" / "к дополнительному соглашению №57"
    for para in paragraphs[:5]:
        if result["appendix_num"] is None:
            appendix_match = re.search(r"[Пп]риложение\s*№\s*(\d+)", para)
            if appendix_match:
                result["appendix_num"] = appendix_match.group(1)
        if result["ds_num"] is None:
            ds_match = re.search(r"соглашени\w*\s*№\s*(\d+)", para, re.IGNORECASE)
            if ds_match:
                result["ds_num"] = ds_match.group(1)
        if result["appendix_num"] and result["ds_num"]:
            break

    # Ищем номер маршрута
    for para in paragraphs[:10]:
        route_match = re.search(r"[Пп]араметры\s+маршрута\s*№\s*(\d+[А-ЯЁа-яёA-Za-z]*)", para)
        if route_match:
            result["route"] = route_match.group(1)
            break
        route_match2 = re.search(r"маршрута\s*№\s*(\d+[А-ЯЁа-яёA-Za-z]*)", para)
        if route_match2:
            result["route"] = route_match2.group(1)
            break

    # Ищем даты
    for para in paragraphs[:10]:
        dates = parse_dates(para)
        if dates["date_from"] or dates["date_to"] or dates["date_on"]:
            result["date_from"] = dates["date_from"]
            result["date_to"] = dates["date_to"]
            result["date_on"] = dates["date_on"]
            break

    # Ищем протяженности
    text = "\n".join(paragraphs)

    sum_match = re.search(
        r"протяж[её]нност[ьъ][^.]*?всего[:\s]+(\d+)[,.](\d+)\s*км",
        text,
        re.IGNORECASE
    )
    if sum_match:
        result["length_sum"] = float(f"{sum_match.group(1)}.{sum_match.group(2)}")

    np_match = re.search(
        r"в\s+прямом\s+направлении[:\s]+(\d+)[,.](\d+)\s*км",
        text,
        re.IGNORECASE
    )
    if np_match:
        result["length_forward"] = float(f"{np_match.group(1)}.{np_match.group(2)}")

    op_match = re.search(
        r"в\s+обратном\s+направлении[:\s]+(\d+)[,.](\d+)\s*км",
        text,
        re.IGNORECASE
    )
    if op_match:
        result["length_reverse"] = float(f"{op_match.group(1)}.{op_match.group(2)}")

    # Извлекаем данные о рейсах из таблиц
    trips_info = extract_trips_from_tables(tables, paragraphs)
    result.update(trips_info)

    return result


def _find_appendix_info_in_texts(texts: list[str]) -> tuple:
    """
    Ищет номер приложения и номер ДС в списке строк текста.

    Примеры входных строк:
      «Приложение №12»
      «к дополнительному соглашению №53»
      «к Государственному контракту №01722000025210002220001...»

    Returns:
        (appendix_number, ds_number) — строки или None
    """
    appendix_number = None
    ds_number = None
    for text in texts:
        app_match = re.search(r'приложение\s*№\s*(\d+)', text, re.IGNORECASE)
        if app_match and appendix_number is None:
            appendix_number = app_match.group(1)
        ds_match = re.search(r'соглашению\s*№\s*(\d+)', text, re.IGNORECASE)
        if ds_match and ds_number is None:
            ds_number = ds_match.group(1)
        if appendix_number and ds_number:
            break
    return appendix_number, ds_number


def extract_raw_tables_for_json(tables: list[list[list[str]]], body_seq: list[tuple] | None = None) -> dict:
    """
    Идентифицирует и возвращает сырые таблицы из основного документа ДС для сохранения в JSON.

    Возвращает словарь с ключами:
    - table_raschet_izm_objema  — «Расчет изменения объема»
    - table_etapy_sroki         — «Этапы исполнения Контракта (в части сроков выполнения работ)»
    - table_finansirovanie_po_godam — «Год / Финансирование по годам, руб.»
    - table_etapy_avans         — «Этапы исполнения Контракта (с учетом порядка погашения авансов)»
    - table_objemy_rabot        — «Объемы работ» (только для ГК252)

    Для каждой таблицы также добавляются ключи *_appendix_number и *_ds_number
    с номером приложения и ДС из заголовочных строк таблицы.
    """
    result = {
        "table_raschet_izm_objema": None,
        "table_raschet_izm_objema_appendix_number": None,
        "table_raschet_izm_objema_ds_number": None,
        "table_etapy_sroki": None,
        "table_etapy_sroki_appendix_number": None,
        "table_etapy_sroki_ds_number": None,
        "table_finansirovanie_po_godam": None,
        "table_finansirovanie_po_godam_appendix_number": None,
        "table_finansirovanie_po_godam_ds_number": None,
        "table_etapy_avans": None,
        "table_etapy_avans_appendix_number": None,
        "table_etapy_avans_ds_number": None,
        "table_objemy_rabot": None,
    }

    # Строим карту: table_index → параграфы непосредственно перед этой таблицей
    # (только параграфы между предыдущей таблицей и текущей)
    table_preceding_paras: dict[int, list[str]] = {}
    if body_seq:
        current_paras: list[str] = []
        for item_type, item_data in body_seq:
            if item_type == 'para':
                current_paras.append(item_data)
            elif item_type == 'table':
                table_preceding_paras[item_data] = current_paras[:]
                current_paras = []

    def get_appendix_info(table_idx: int) -> tuple:
        """Ищет номер приложения и ДС в параграфах перед таблицей table_idx."""
        paras = table_preceding_paras.get(table_idx, [])
        return _find_appendix_info_in_texts(list(reversed(paras)))

    for t_idx, table in enumerate(tables):
        if not table or len(table) < 2:
            continue

        # Объединяем текст первых 3 строк для идентификации заголовка
        header_text = " ".join(
            " ".join(cell.lower() for cell in row)
            for row in table[:3]
        )

        # «Расчет изменения объема»: содержит «предлагаемые изменения» в заголовке
        if result["table_raschet_izm_objema"] is None and "предлагаемые изменения" in header_text:
            result["table_raschet_izm_objema"] = table
            app_num, ds_num = get_appendix_info(t_idx)
            result["table_raschet_izm_objema_appendix_number"] = app_num
            result["table_raschet_izm_objema_ds_number"] = ds_num
            continue

        # «Финансирование по годам»: содержит «финансирование по годам» в заголовке
        if result["table_finansirovanie_po_godam"] is None and "финансирование по годам" in header_text:
            result["table_finansirovanie_po_godam"] = table
            app_num, ds_num = get_appendix_info(t_idx)
            result["table_finansirovanie_po_godam_appendix_number"] = app_num
            result["table_finansirovanie_po_godam_ds_number"] = ds_num
            continue

        # Таблицы этапов — содержат «максимальная транспортная работа»
        if "максимальная транспортная работа" in header_text:
            has_price = "стоимость" in header_text
            has_avans = "авансов" in header_text
            has_raschet_period = "расчетный период" in header_text

            # «Этапы с учетом авансов»: есть столбец «авансов»
            if result["table_etapy_avans"] is None and has_avans:
                result["table_etapy_avans"] = table
                app_num, ds_num = get_appendix_info(t_idx)
                result["table_etapy_avans_appendix_number"] = app_num
                result["table_etapy_avans_ds_number"] = ds_num
                continue

            # «Объемы работ» (ГК252): нет столбца цены или есть «расчетный период»
            if result["table_objemy_rabot"] is None and (not has_price or has_raschet_period):
                result["table_objemy_rabot"] = table
                continue

            # «Этапы в части сроков»: есть столбец цены, нет авансов
            if result["table_etapy_sroki"] is None and has_price:
                result["table_etapy_sroki"] = table
                app_num, ds_num = get_appendix_info(t_idx)
                result["table_etapy_sroki_appendix_number"] = app_num
                result["table_etapy_sroki_ds_number"] = ds_num
                continue

    # Финансирование по годам идёт в том же приложении, что и этапы (сроки)
    if result["table_finansirovanie_po_godam_appendix_number"] is None:
        result["table_finansirovanie_po_godam_appendix_number"] = result["table_etapy_sroki_appendix_number"]
        result["table_finansirovanie_po_godam_ds_number"] = result["table_etapy_sroki_ds_number"]

    return result


def process_zip(archive_path: Path, verbose: bool = True) -> dict:
    """Обрабатывает ZIP архив."""
    with zipfile.ZipFile(archive_path, "r") as zf:
        files = zf.namelist()
        main_doc = find_main_document(files)

        if not main_doc:
            raise ValueError("Основной документ ДС не найден в архиве.")

        # Обрабатываем основной документ
        if verbose:
            print("Чтение основного документа...", end=" ", flush=True)
        docx_bytes = zf.read(main_doc)
        if verbose:
            print("OK")
            print("  Парсинг docx...", end=" ", flush=True)
        paragraphs, tables, body_seq = get_docx_data(docx_bytes)
        if verbose:
            print("OK")
            print("  Извлечение общей информации...", end=" ", flush=True)
        general = extract_general_info(paragraphs)
        if verbose:
            print("OK")
            print("  Извлечение значений из таблиц...", end=" ", flush=True)
        general.update(extract_values_from_tables(tables))
        if verbose:
            print("OK")
            print("  Извлечение изменений...", end=" ", flush=True)
        change_with_money, change_with_money_no_appendix, change_without_money, change_without_money_no_appendix = extract_changes(paragraphs, general.get("contract_number"))
        if verbose:
            print("OK")

        # Обрабатываем папки приложений
        appendix_folders = find_appendix_folders(files)
        appendices = {}
        total_appendices = len(appendix_folders)

        # Собираем все изменения для определения маршрутов
        all_changes = change_with_money + change_without_money

        for idx, (appendix_num, folder_files) in enumerate(sorted(appendix_folders.items(), key=lambda x: int(x[0])), 1):
            # Определяем маршрут для этого приложения
            route = get_route_for_appendix(appendix_num, all_changes)

            # Специальная обработка для маршрута 315 (XLS формат)
            if route == "315" and ROUTE_315_SUPPORT:
                xls_file = find_xls_file_in_folder(folder_files)
                if xls_file:
                    if verbose:
                        print(f"Приложение {appendix_num} (маршрут 315, XLS)...", end=" ", flush=True)
                    try:
                        # Извлекаем XLS во временный файл
                        xls_bytes = zf.read(xls_file)
                        with tempfile.NamedTemporaryFile(suffix=".xls", delete=False) as tmp:
                            tmp.write(xls_bytes)
                            tmp_path = tmp.name

                        # Парсим XLS
                        result_315 = parse_route_315_xls(tmp_path)

                        # Добавляем оба результата в appendices
                        appendices[appendix_num] = result_315["route_315"]
                        appendices[f"{appendix_num}."] = result_315["route_315_dot"]

                        # Удаляем временный файл
                        os.unlink(tmp_path)

                        if verbose:
                            print("OK")
                        continue
                    except Exception as e:
                        if verbose:
                            print(f"ОШИБКА: {e}")
                        appendices[appendix_num] = {"error": str(e)}
                        continue

            # Стандартная обработка через Word документ
            contract_doc = find_contract_doc_in_folder(folder_files)
            if contract_doc:
                if verbose:
                    print(f"Приложение {appendix_num} ({idx}/{total_appendices})...", end=" ", flush=True)
                try:
                    doc_bytes = zf.read(contract_doc)
                    doc_paragraphs, doc_tables, _ = get_word_data(doc_bytes, contract_doc)
                    appendix_info = extract_appendix_info(doc_paragraphs, doc_tables)
                    appendices[appendix_num] = appendix_info
                    if verbose:
                        print("OK")
                except Exception as e:
                    if verbose:
                        print(f"ОШИБКА: {e}")
                    appendices[appendix_num] = {"error": str(e)}

        # Проверяем наличие высвобождения в документе
        vysvobozhdenie_data = detect_vysvobozhdenie(paragraphs, tables)
        
        result = {
            "general": general,
            "change_with_money": change_with_money,
            "change_with_money_no_appendix": change_with_money_no_appendix,
            "change_without_money": change_without_money,
            "change_without_money_no_appendix": change_without_money_no_appendix,
            "appendices": appendices
        }

        # Добавляем данные о высвобождении если найдены
        if vysvobozhdenie_data:
            result["vysvobozhdenie"] = vysvobozhdenie_data

        km_data = extract_stages_km_252(tables)
        if km_data:
            km_data["ds_number"] = general.get("ds_number")
            result["km_data"] = km_data

        result.update(extract_raw_tables_for_json(tables, body_seq))

        return result


class ExtractedArchive:
    """Обёртка для работы с распакованным архивом как с ZIP/RAR."""

    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self._files = None

    def namelist(self) -> list[str]:
        """Возвращает список файлов в архиве."""
        if self._files is None:
            self._files = []
            for f in self.base_dir.rglob("*"):
                if f.is_file():
                    rel = f.relative_to(self.base_dir)
                    self._files.append(str(rel).replace("\\", "/"))
        return self._files

    def read(self, filename: str) -> bytes:
        """Читает файл из архива."""
        file_path = self.base_dir / filename.replace("/", os.sep)
        if not file_path.exists():
            # Попробуем найти файл без учёта регистра
            parts = filename.replace("/", os.sep).split(os.sep)
            current = self.base_dir
            for part in parts:
                found = None
                for child in current.iterdir():
                    if child.name.lower() == part.lower():
                        found = child
                        break
                if found:
                    current = found
                else:
                    raise FileNotFoundError(f"Файл не найден: {filename}")
            file_path = current
        return file_path.read_bytes()


def process_rar(archive_path: Path, verbose: bool = True) -> dict:
    """Обрабатывает RAR архив."""
    # Сначала пробуем распаковать через 7z (более надёжно)
    sevenzip = find_7z()

    if sevenzip:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            if extract_rar_with_7z(archive_path, tmpdir_path):
                # Используем обёртку для работы с распакованными файлами
                rf = ExtractedArchive(tmpdir_path)
                return _process_rar_contents(rf, verbose)

    # Fallback на rarfile
    if not RAR_SUPPORT:
        raise ImportError("Для работы с RAR установите 7-Zip или pip install rarfile")

    with rarfile.RarFile(archive_path, "r") as rf:
        return _process_rar_contents(rf, verbose)


def _process_rar_contents(rf, verbose: bool = True) -> dict:
    """Обрабатывает содержимое RAR архива."""
    files = rf.namelist()
    main_doc = find_main_document(files)

    if not main_doc:
        raise ValueError("Основной документ ДС не найден в архиве.")

    if verbose:
        print("Чтение основного документа...", end=" ", flush=True)
    docx_bytes = rf.read(main_doc)
    if verbose:
        print("OK")
        print("  Парсинг docx...", end=" ", flush=True)
    paragraphs, tables, body_seq = get_docx_data(docx_bytes)
    if verbose:
        print("OK")
        print("  Извлечение общей информации...", end=" ", flush=True)
    general = extract_general_info(paragraphs)
    if verbose:
        print("OK")
        print("  Извлечение значений из таблиц...", end=" ", flush=True)
    general.update(extract_values_from_tables(tables))
    if verbose:
        print("OK")
        print("  Извлечение изменений...", end=" ", flush=True)
    change_with_money, change_with_money_no_appendix, change_without_money, change_without_money_no_appendix = extract_changes(paragraphs, general.get("contract_number"))
    if verbose:
        print("OK")

    appendix_folders = find_appendix_folders(files)
    appendices = {}
    total_appendices = len(appendix_folders)

    # Собираем все изменения для определения маршрутов
    all_changes = change_with_money + change_without_money

    for idx, (appendix_num, folder_files) in enumerate(sorted(appendix_folders.items(), key=lambda x: int(x[0])), 1):
        # Определяем маршрут для этого приложения
        route = get_route_for_appendix(appendix_num, all_changes)

        # Специальная обработка для маршрута 315 (XLS формат)
        if route == "315" and ROUTE_315_SUPPORT:
            xls_file = find_xls_file_in_folder(folder_files)
            if xls_file:
                if verbose:
                    print(f"Приложение {appendix_num} (маршрут 315, XLS)...", end=" ", flush=True)
                try:
                    # Извлекаем XLS во временный файл
                    xls_bytes = rf.read(xls_file)
                    with tempfile.NamedTemporaryFile(suffix=".xls", delete=False) as tmp:
                        tmp.write(xls_bytes)
                        tmp_path = tmp.name

                    # Парсим XLS
                    result_315 = parse_route_315_xls(tmp_path)

                    # Добавляем оба результата в appendices
                    appendices[appendix_num] = result_315["route_315"]
                    appendices[f"{appendix_num}."] = result_315["route_315_dot"]

                    # Удаляем временный файл
                    os.unlink(tmp_path)

                    if verbose:
                        print("OK")
                    continue
                except Exception as e:
                    if verbose:
                        print(f"ОШИБКА: {e}")
                    appendices[appendix_num] = {"error": str(e)}
                    continue

        # Стандартная обработка через Word документ
        contract_doc = find_contract_doc_in_folder(folder_files)
        if contract_doc:
            if verbose:
                print(f"Приложение {appendix_num} ({idx}/{total_appendices})...", end=" ", flush=True)
            try:
                doc_bytes = rf.read(contract_doc)
                doc_paragraphs, doc_tables, _ = get_word_data(doc_bytes, contract_doc)
                appendix_info = extract_appendix_info(doc_paragraphs, doc_tables)
                appendices[appendix_num] = appendix_info
                if verbose:
                    print("OK")
            except Exception as e:
                if verbose:
                    print(f"ОШИБКА: {e}")
                appendices[appendix_num] = {"error": str(e)}

    # Проверяем наличие высвобождения в документе
    vysvobozhdenie_data = detect_vysvobozhdenie(paragraphs, tables)

    result = {
        "general": general,
        "change_with_money": change_with_money,
        "change_with_money_no_appendix": change_with_money_no_appendix,
        "change_without_money": change_without_money,
        "change_without_money_no_appendix": change_without_money_no_appendix,
        "appendices": appendices
    }

    # Добавляем данные о высвобождении если найдены
    if vysvobozhdenie_data:
        result["vysvobozhdenie"] = vysvobozhdenie_data

    km_data = extract_stages_km_252(tables)
    if km_data:
        km_data["ds_number"] = general.get("ds_number")
        result["km_data"] = km_data

    result.update(extract_raw_tables_for_json(tables, body_seq))

    return result


def process_archive(archive_path: str | Path, verbose: bool = True) -> dict:
    """Обрабатывает архив (ZIP или RAR)."""
    path = Path(archive_path)

    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {path}")

    suffix = path.suffix.lower()

    if suffix == ".zip":
        return process_zip(path, verbose=verbose)
    elif suffix == ".rar":
        return process_rar(path, verbose=verbose)
    else:
        raise ValueError(f"Неподдерживаемый формат архива: {suffix}")


def main():
    if len(sys.argv) < 2:
        script_dir = Path(__file__).parent
        archives = list(script_dir.glob("*.zip")) + list(script_dir.glob("*.rar"))

        if not archives:
            print("Использование: python extract_contract_info.py <путь_к_архиву>")
            print("Или поместите архив в ту же папку, что и скрипт")
            sys.exit(1)

        archive_path = archives[0]
        print(f"Найден архив: {archive_path.name}")
    else:
        archive_path = Path(sys.argv[1])

    try:
        info = process_archive(archive_path)

        output_path = Path(archive_path).with_suffix(".json")

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(info, f, ensure_ascii=False, indent=2)

        print(f"Результат сохранен в: {output_path}")

    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
