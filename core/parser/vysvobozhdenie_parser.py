"""
Парсер ДС на высвобождение (doc/docx).

Извлекает:
- Номер ДС и номер ГК из заголовка
- Номер закрываемого этапа и сумму факта из п.1
- Новую цену контракта из п.2
- Таблицу 1: «Этапы исполнения Контракта (в части сроков выполнения работ)»
- Таблицу 2: «Этапы исполнения Контракта (с учётом порядка погашения авансов)»
"""

import re
from io import BytesIO
from pathlib import Path

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import win32com.client
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False


# ---------------------------------------------------------------------------
# Вспомогательные функции
# ---------------------------------------------------------------------------

def _parse_money_amount(text: str) -> float | None:
    """Парсит сумму в рублях из текста. Возвращает float или None."""
    if not text:
        return None
    # «NNN рублей YY копеек» (с пробелами/неразрывными пробелами внутри числа)
    match = re.search(
        r"([\d\s\u00a0]+)\s*руб\w*\s+(\d+)\s*коп\w*",
        text,
        re.IGNORECASE,
    )
    if match:
        rubles = re.sub(r"[\s\u00a0]", "", match.group(1))
        kopeks = match.group(2)
        try:
            return float(f"{rubles}.{kopeks}")
        except ValueError:
            pass

    # «NNN рублей» без копеек
    match = re.search(r"([\d\s\u00a0]+)\s*руб\w*", text, re.IGNORECASE)
    if match:
        rubles = re.sub(r"[\s\u00a0]", "", match.group(1))
        try:
            return float(rubles)
        except ValueError:
            pass

    return None


def _parse_number(text: str) -> float | None:
    """Парсит число из ячейки таблицы (1 234,56 → 1234.56)."""
    if not text:
        return None
    cleaned = re.sub(r"[\s\u00a0]", "", text).replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Чтение документа
# ---------------------------------------------------------------------------

def _get_docx_data(file_bytes: bytes) -> tuple[list[str], list[list[list[str]]]]:
    """Извлекает параграфы и таблицы из .docx файла."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx не установлен")

    doc = Document(BytesIO(file_bytes))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip().replace("\n", " "))
            rows.append(cells)
        tables.append(rows)

    return paragraphs, tables


def _get_doc_data_win32(file_path: Path) -> tuple[list[str], list[list[list[str]]]]:
    """Извлекает параграфы и таблицы из .doc файла через win32com."""
    if not DOC_SUPPORT:
        raise RuntimeError("win32com не установлен — .doc файлы не поддерживаются")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(str(file_path.resolve()))

        paragraphs = []
        for para in doc.Paragraphs:
            text = para.Range.Text.strip().rstrip("\r\x07")
            if text:
                paragraphs.append(text)

        tables = []
        for table in doc.Tables:
            rows = []
            for r in range(1, table.Rows.Count + 1):
                cells = []
                for c in range(1, table.Columns.Count + 1):
                    try:
                        text = table.Cell(r, c).Range.Text.strip().rstrip("\r\x07")
                    except Exception:
                        text = ""
                    cells.append(text)
                rows.append(cells)
            tables.append(rows)

        return paragraphs, tables

    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


# ---------------------------------------------------------------------------
# Парсинг содержимого
# ---------------------------------------------------------------------------

def _parse_header(paragraphs: list[str]) -> dict:
    """Извлекает номер ДС и полный номер ГК из заголовочных параграфов."""
    result = {"ds_number": None, "contract_number": None}

    text = "\n".join(paragraphs[:10])  # заголовок — первые параграфы

    # Номер ДС: «Дополнительное соглашение № 43»
    ds_match = re.search(
        r"[Дд]ополнительное\s+соглашение\s*[№#N]\s*(\d+)",
        text,
        re.IGNORECASE,
    )
    if ds_match:
        result["ds_number"] = ds_match.group(1)

    # Полный номер ГК: длинная цифровая строка
    gk_match = re.search(r"[Кк]онтракту?\s*[№#N]\s*(\d{10,})", text, re.IGNORECASE)
    if gk_match:
        result["contract_number"] = gk_match.group(1)

    return result


def _get_contract_short_number(full_number: str | None) -> str | None:
    """Выводит краткий номер контракта из полного (аналог логики upload.py)."""
    if not full_number:
        return None
    # Последние 7 цифр перед «0001» → позиции [-7:-4]
    if len(full_number) >= 7 and full_number.endswith("0001"):
        return full_number[-7:-4]
    return None


def _parse_clause1(paragraphs: list[str]) -> dict:
    """
    Ищет абзац вида:
    «...стоимость фактически выполненных ... за N этап ... составила X рублей Y копеек»
    Возвращает {closed_stage: int, closed_amount: float}.
    """
    result = {"closed_stage": None, "closed_amount": None}

    full_text = " ".join(paragraphs)

    stage_match = re.search(
        r"за\s+(\d+)\s*этап",
        full_text,
        re.IGNORECASE,
    )
    if stage_match:
        result["closed_stage"] = int(stage_match.group(1))

    # Ищем «составила X рублей Y копеек»
    amount_match = re.search(
        r"составила\s+([\d\s\u00a0]+\s*руб\w*(?:\s+\d+\s*коп\w*)?)",
        full_text,
        re.IGNORECASE,
    )
    if amount_match:
        result["closed_amount"] = _parse_money_amount(amount_match.group(1))

    return result


def _parse_clause2(paragraphs: list[str]) -> dict:
    """
    Ищет абзац вида:
    «цена Контракта составит X рублей Y копеек»
    Возвращает {new_contract_price: float}.
    """
    result = {"new_contract_price": None}

    full_text = " ".join(paragraphs)

    price_match = re.search(
        r"цена\s+[Кк]онтракта\s+составит\s+([\d\s\u00a0]+\s*руб\w*(?:\s+\d+\s*коп\w*)?)",
        full_text,
        re.IGNORECASE,
    )
    if price_match:
        result["new_contract_price"] = _parse_money_amount(price_match.group(1))

    return result


def _is_stages_table1_header(header_row: list[str]) -> bool:
    """Определяет, является ли строка заголовком таблицы 1 (сроки выполнения работ)."""
    joined = " ".join(header_row).lower()
    has_km = "транспортная работа" in joined or "максимальная" in joined
    has_price = "стоимость" in joined
    # Таблица 2 имеет доп. столбец «с учётом выплаченных авансов»
    has_avans = "аванс" in joined
    return has_km and has_price and not has_avans


def _is_stages_table2_header(header_row: list[str]) -> bool:
    """Определяет, является ли строка заголовком таблицы 2 (с учётом авансов)."""
    joined = " ".join(header_row).lower()
    has_km = "транспортная работа" in joined or "максимальная" in joined
    has_avans = "аванс" in joined
    return has_km and has_avans


def _is_finansirovanie_header(header_row: list[str]) -> bool:
    """Определяет, является ли строка заголовком таблицы «Финансирование по годам»."""
    joined = " ".join(header_row).lower()
    return "финансирование" in joined and "транспортная" not in joined


def _parse_stage_row1(row: list[str]) -> dict | None:
    """
    Парсит строку таблицы 1.
    Столбцы: №, Год, Квартал/месяц, Срок начала-окончания, Км, Цена руб.
    Возвращает dict или None если строка пустая / ИТОГО.
    """
    if not row or len(row) < 4:
        return None

    num_raw = row[0].strip()
    if not num_raw:
        return None

    is_total = num_raw.upper().startswith("ИТОГО")

    if is_total:
        km = _parse_number(row[4]) if len(row) > 4 else None
        price = _parse_number(row[5]) if len(row) > 5 else None
        return {
            "stage": None,
            "year": None,
            "period": None,
            "date_range": None,
            "km": km,
            "price": price,
            "is_total": True,
        }

    try:
        stage = int(num_raw)
    except ValueError:
        return None

    year_raw = row[1].strip() if len(row) > 1 else ""
    try:
        year = int(year_raw)
    except ValueError:
        year = None

    period = row[2].strip() if len(row) > 2 else ""
    date_range = row[3].strip() if len(row) > 3 else ""
    km = _parse_number(row[4]) if len(row) > 4 else None
    price = _parse_number(row[5]) if len(row) > 5 else None

    return {
        "stage": stage,
        "year": year,
        "period": period,
        "date_range": date_range,
        "km": km,
        "price": price,
        "is_total": False,
    }


def _parse_stage_row2(row: list[str]) -> dict | None:
    """
    Парсит строку таблицы 2.
    Столбцы: №, Год, Квартал/месяц, Срок начала-окончания,
             Км, Цена руб., Цена с учётом авансов, Сроки оплаты.
    """
    if not row or len(row) < 4:
        return None

    num_raw = row[0].strip()
    if not num_raw:
        return None

    is_total = num_raw.upper().startswith("ИТОГО")

    if is_total:
        km = _parse_number(row[4]) if len(row) > 4 else None
        price = _parse_number(row[5]) if len(row) > 5 else None
        price_avans = _parse_number(row[6]) if len(row) > 6 else None
        return {
            "stage": None,
            "year": None,
            "period": None,
            "date_range": None,
            "km": km,
            "price": price,
            "price_avans": price_avans,
            "payment_dates": None,
            "is_total": True,
        }

    try:
        stage = int(num_raw)
    except ValueError:
        return None

    year_raw = row[1].strip() if len(row) > 1 else ""
    try:
        year = int(year_raw)
    except ValueError:
        year = None

    period = row[2].strip() if len(row) > 2 else ""
    date_range = row[3].strip() if len(row) > 3 else ""
    km = _parse_number(row[4]) if len(row) > 4 else None
    price = _parse_number(row[5]) if len(row) > 5 else None
    price_avans = _parse_number(row[6]) if len(row) > 6 else None
    payment_dates = row[7].strip() if len(row) > 7 else ""

    return {
        "stage": stage,
        "year": year,
        "period": period,
        "date_range": date_range,
        "km": km,
        "price": price,
        "price_avans": price_avans,
        "payment_dates": payment_dates or None,
        "is_total": False,
    }


def _find_and_parse_tables(
    tables: list[list[list[str]]],
) -> tuple[list[dict], list[dict], list[list[str]], list[list[str]], list[list[str]]]:
    """
    Находит таблицы 1, 2 и «Финансирование по годам» среди всех таблиц документа и парсит их.
    Возвращает (rows_table1, rows_table2, raw_table1, raw_table2, raw_finansirovanie).
    raw_table* — оригинальные строки ячеек (включая заголовок и ИТОГО),
    нужны для совместимости с функциями проверки вкладки «Проверка таблиц».
    """
    rows1: list[dict] = []
    rows2: list[dict] = []
    raw1: list[list[str]] = []
    raw2: list[list[str]] = []
    raw_fin: list[list[str]] = []

    for table in tables:
        if not table:
            continue

        header = table[0]

        # Таблица финансирования может иметь мало столбцов — проверяем её первой
        if not raw_fin and _is_finansirovanie_header(header):
            raw_fin = table
            continue

        if len(header) < 4:
            continue

        if _is_stages_table2_header(header):
            raw2 = table
            for row in table[1:]:
                parsed = _parse_stage_row2(row)
                if parsed is not None:
                    rows2.append(parsed)
        elif _is_stages_table1_header(header):
            raw1 = table
            for row in table[1:]:
                parsed = _parse_stage_row1(row)
                if parsed is not None:
                    rows1.append(parsed)

    return rows1, rows2, raw1, raw2, raw_fin


# ---------------------------------------------------------------------------
# Главная точка входа
# ---------------------------------------------------------------------------

def parse_vysvobozhdenie(file_bytes: bytes, filename: str) -> dict:
    """
    Парсит документ ДС на высвобождение (.doc/.docx).

    Args:
        file_bytes: байты файла
        filename: имя файла (для определения расширения)

    Returns:
        dict с ключами:
          type, general, stages_table1, stages_table2,
          itogo_km, itogo_price
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        paragraphs, tables = _get_docx_data(file_bytes)
    elif ext == ".doc":
        # Для .doc нужен временный файл
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = Path(tmp.name)
        try:
            paragraphs, tables = _get_doc_data_win32(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    header = _parse_header(paragraphs)
    clause1 = _parse_clause1(paragraphs)
    clause2 = _parse_clause2(paragraphs)

    contract_short = _get_contract_short_number(header.get("contract_number"))

    general = {
        "ds_number": header.get("ds_number"),
        "contract_number": header.get("contract_number"),
        "contract_short_number": contract_short,
        "closed_stage": clause1.get("closed_stage"),
        "closed_amount": clause1.get("closed_amount"),
        "new_contract_price": clause2.get("new_contract_price"),
    }

    stages_table1, stages_table2, raw_table1, raw_table2, raw_finansirovanie = _find_and_parse_tables(tables)

    # Извлекаем строку ИТОГО из таблицы 1
    itogo_km: float | None = None
    itogo_price: float | None = None
    for row in stages_table1:
        if row.get("is_total"):
            itogo_km = row.get("km")
            itogo_price = row.get("price")
            break

    # Извлекаем строку ИТОГО из таблицы 2
    itogo_km_t2: float | None = None
    itogo_price_t2: float | None = None
    for row in stages_table2:
        if row.get("is_total"):
            itogo_km_t2 = row.get("km")
            itogo_price_t2 = row.get("price")
            break

    return {
        "type": "vysvobozhdenie",
        "general": general,
        "stages_table1": [r for r in stages_table1 if not r.get("is_total")],
        "stages_table2": [r for r in stages_table2 if not r.get("is_total")],
        # Сырые таблицы (список строк, каждая строка — список ячеек в виде строк)
        # используются вкладкой «Проверка таблиц» (аналог table_etapy_sroki / table_etapy_avans)
        "stages_table1_raw": raw_table1,
        "stages_table2_raw": raw_table2,
        "stages_finansirovanie_raw": raw_finansirovanie,
        "itogo_km": itogo_km,
        "itogo_price": itogo_price,
        "itogo_km_t2": itogo_km_t2,
        "itogo_price_t2": itogo_price_t2,
    }
