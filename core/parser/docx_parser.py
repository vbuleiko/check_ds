"""Парсинг DOCX документов."""
import re
from io import BytesIO
from pathlib import Path

from docx import Document


def get_docx_data(file_source: str | Path | bytes) -> tuple[list[str], list[list[list[str]]]]:
    """
    Извлекает параграфы и таблицы из docx файла.

    Args:
        file_source: Путь к файлу или байты

    Returns:
        (paragraphs, tables) где tables - список таблиц,
        каждая таблица - список строк, каждая строка - список ячеек
    """
    if isinstance(file_source, bytes):
        doc = Document(BytesIO(file_source))
    else:
        doc = Document(file_source)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables = []
    for table in doc.tables:
        table_data = []
        # Отслеживаем _tc элементы по индексу колонки для определения вертикальных объединений
        seen_tcs: dict[int, set] = {}
        for row in table.rows:
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                if col_idx not in seen_tcs:
                    seen_tcs[col_idx] = set()
                if tc in seen_tcs[col_idx]:
                    # Ячейка является продолжением вертикального объединения — возвращаем пустую строку
                    row_data.append("")
                else:
                    seen_tcs[col_idx].add(tc)
                    row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables.append(table_data)

    return paragraphs, tables


def extract_number(text: str) -> float | int | None:
    """Извлекает число из текста."""
    if not text:
        return None

    # Убираем пробелы и заменяем запятую на точку
    text = text.replace(" ", "").replace("\u00a0", "").replace(",", ".")

    # Ищем число
    match = re.search(r"-?\d+\.?\d*", text)
    if match:
        try:
            val = float(match.group())
            if val == int(val):
                return int(val)
            return round(val, 2)
        except ValueError:
            return None
    return None


def extract_date(text: str) -> str | None:
    """
    Извлекает дату из текста.

    Поддерживает форматы:
    - DD.MM.YYYY
    - DD.MM.YY
    - YYYY-MM-DD

    Returns:
        Дата в формате YYYY-MM-DD или None
    """
    if not text:
        return None

    # DD.MM.YYYY
    match = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", text)
    if match:
        d, m, y = match.groups()
        return f"{y}-{m}-{d}"

    # DD.MM.YY
    match = re.search(r"(\d{2})\.(\d{2})\.(\d{2})", text)
    if match:
        d, m, y = match.groups()
        year = f"20{y}" if int(y) < 50 else f"19{y}"
        return f"{year}-{m}-{d}"

    # YYYY-MM-DD
    match = re.search(r"(\d{4})-(\d{2})-(\d{2})", text)
    if match:
        return match.group()

    return None


def extract_date_range(text: str) -> tuple[str | None, str | None, str | None]:
    """
    Извлекает диапазон дат из текста.

    Форматы:
    - "с DD.MM.YYYY года по DD.MM.YYYY года" → (date_from, date_to, None)
    - "с DD.MM.YYYY года" → (date_from, None, None)
    - "на DD.MM.YYYY года" → (None, None, date_on)
    - "DD.MM.YYYY года" (в начале подпункта) → (None, None, date_on)

    Returns:
        (date_from, date_to, date_on) в формате YYYY-MM-DD
    """
    if not text:
        return None, None, None

    text = text.lower()

    # Диапазон: "с DD.MM.YYYY года по DD.MM.YYYY года"
    match = re.search(
        r"с\s+(\d{2})\.(\d{2})\.(\d{4})\s+года\s+по\s+(\d{2})\.(\d{2})\.(\d{4})\s+года",
        text
    )
    if match:
        d1, m1, y1, d2, m2, y2 = match.groups()
        return f"{y1}-{m1}-{d1}", f"{y2}-{m2}-{d2}", None

    # Начало: "с DD.MM.YYYY года"
    match = re.search(r"с\s+(\d{2})\.(\d{2})\.(\d{4})\s+года", text)
    if match:
        d, m, y = match.groups()
        return f"{y}-{m}-{d}", None, None

    # Конкретная дата: "на DD.MM.YYYY года"
    match = re.search(r"на\s+(\d{2})\.(\d{2})\.(\d{4})\s+года", text)
    if match:
        d, m, y = match.groups()
        return None, None, f"{y}-{m}-{d}"

    # Конкретная дата в начале строки: "DD.MM.YYYY года"
    # Учитывает возможный номер подпункта в начале (например, "1.3. 16.11.2025 года")
    match = re.search(r"^\d+\.\d+\.\s*(\d{2})\.(\d{2})\.(\d{4})\s+года", text)
    if match:
        d, m, y = match.groups()
        return None, None, f"{y}-{m}-{d}"
    
    # Конкретная дата в начале строки без номера подпункта
    match = re.match(r"^(\d{2})\.(\d{2})\.(\d{4})\s+года", text)
    if match:
        d, m, y = match.groups()
        return None, None, f"{y}-{m}-{d}"

    return None, None, None


def find_table_by_header(
    tables: list[list[list[str]]],
    header_keywords: list[str]
) -> list[list[str]] | None:
    """
    Находит таблицу по ключевым словам в заголовке.

    Args:
        tables: Список таблиц
        header_keywords: Ключевые слова для поиска

    Returns:
        Найденная таблица или None
    """
    for table in tables:
        if not table:
            continue

        # Проверяем первые строки таблицы
        header_text = " ".join(" ".join(row) for row in table[:3]).lower()

        if all(kw.lower() in header_text for kw in header_keywords):
            return table

    return None
